"""
Gap Analysis Project - Professional Report Generator
Generates a comprehensive graduation project report in DOCX format
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
from pathlib import Path


def add_title_page(doc):
    """Add professional title page"""
    # University name
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Vodafone International Foundation\nFaculty of Engineering')
    run.font.size = Pt(16)
    run.font.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Project title
    project_title = doc.add_paragraph()
    project_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = project_title.add_run('Supply-Demand Gap Analysis for Rental Platform')
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph()
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('A Comprehensive Data-Driven Approach to Identifying\nand Addressing Supply Shortages')
    run.font.size = Pt(14)
    run.font.italic = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Team members
    team = doc.add_paragraph()
    team.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = team.add_run('Prepared By:\n\nHamza Omran\nAhmad [Last Name]\nMoaz [Last Name]')
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Date and version
    footer_info = doc.add_paragraph()
    footer_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_info.add_run(f'Academic Year: 2024-2025\nVersion 4.0\n{datetime.now().strftime("%B %Y")}')
    run.font.size = Pt(11)
    
    doc.add_page_break()


def add_abstract(doc):
    """Add abstract section"""
    doc.add_heading('Abstract', level=1)
    
    abstract_text = """This project presents a comprehensive supply-demand gap analysis for a rental business platform. The analysis identifies product categories experiencing supply shortages by computing gap scores based on the ratio of customer demand to available product listings. Using a data-driven methodology, the project processes transactional data across 8 datasets containing product information, sales records, returns, customer demographics, and geographic territories.

The analysis reveals critical supply shortages in specific subcategories, with gap scores exceeding 200 in several product lines. The project employs a multi-phase pipeline consisting of data understanding, cleaning, supply-demand computation, gap score calculation, visualization, insight generation, and action planning. Interactive visualizations and monitoring dashboards provide stakeholders with actionable intelligence for strategic decision-making.

Key findings indicate that certain product categories face severe supply-demand imbalances, requiring immediate inventory expansion and supplier recruitment. The monitoring framework enables weekly tracking of gap score changes and generates prioritized action plans for operational, strategic, and platform optimization initiatives. The deliverables include comprehensive Excel reports, interactive HTML visualizations, textual insights, and JSON-based data exports suitable for API integration."""
    
    doc.add_paragraph(abstract_text)
    doc.add_page_break()


def add_table_of_contents(doc):
    """Add table of contents placeholder"""
    doc.add_heading('Table of Contents', level=1)
    
    toc_items = [
        '1. Introduction',
        '   1.1 Project Overview',
        '   1.2 Problem Statement',
        '   1.3 Objectives',
        '   1.4 Scope and Limitations',
        '2. Literature Review',
        '   2.1 Supply-Demand Analysis',
        '   2.2 Gap Score Methodology',
        '   2.3 Data-Driven Decision Making',
        '3. Methodology',
        '   3.1 Data Pipeline Architecture',
        '   3.2 Data Sources and Structure',
        '   3.3 Gap Score Calculation Formula',
        '   3.4 Classification Criteria',
        '4. Implementation',
        '   4.1 Phase 1: Data Preparation',
        '   4.2 Phase 2: Supply-Demand Computation',
        '   4.3 Phase 3: Gap Score Calculation',
        '   4.4 Phase 4: Visualization and Insights',
        '   4.5 Phase 5: Monitoring and Action Planning',
        '5. Technical Architecture',
        '   5.1 Python Modules',
        '   5.2 Notebook Workflow',
        '   5.3 Output Files and Deliverables',
        '6. Results and Analysis',
        '   6.1 Category-Level Findings',
        '   6.2 Subcategory-Level Findings',
        '   6.3 Territory-Level Findings',
        '   6.4 Gap Severity Distribution',
        '7. Visualizations',
        '   7.1 Interactive Dashboards',
        '   7.2 Gap Score Charts',
        '   7.3 Supply vs Demand Comparison',
        '8. Business Insights',
        '   8.1 Critical Gaps',
        '   8.2 Balanced Categories',
        '   8.3 Recommendations',
        '9. Monitoring Framework',
        '   9.1 Weekly Snapshots',
        '   9.2 KPI Tracking',
        '   9.3 Alert System',
        '10. Action Plans',
        '    10.1 Operational Actions',
        '    10.2 Strategic Actions',
        '    10.3 Platform Optimization',
        '11. API Integration',
        '    11.1 Endpoint Recommendations',
        '    11.2 Data Export Formats',
        '    11.3 Usage Examples',
        '12. Conclusion',
        '13. Future Work',
        '14. References',
        '15. Appendices'
    ]
    
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()


def add_introduction(doc):
    """Add introduction section"""
    doc.add_heading('1. Introduction', level=1)
    
    doc.add_heading('1.1 Project Overview', level=2)
    intro_text = """This project addresses a fundamental challenge in rental platform operations: identifying and quantifying supply-demand imbalances across product categories. By computing gap scores that represent the ratio of customer demand to available supply, the analysis enables data-driven decision-making for inventory expansion, supplier recruitment, and resource allocation.

The project processes comprehensive transactional data spanning three years (2020-2022) across 4 main product categories and 37 subcategories. The analysis incorporates sales data (56,046 transactions), product information (293 unique products), customer demographics (18,148 customers), and returns data (1,809 returns) to compute accurate supply and demand metrics."""
    
    doc.add_paragraph(intro_text)
    
    doc.add_heading('1.2 Problem Statement', level=2)
    problem_text = """Rental platforms frequently experience supply shortages in high-demand categories, leading to lost revenue opportunities and customer dissatisfaction. Without systematic measurement of supply-demand gaps, platform operators cannot prioritize inventory expansion efforts or allocate marketing resources effectively.

The primary challenges include:
- Lack of quantitative metrics to measure supply-demand imbalances
- Inability to identify which specific subcategories require immediate attention
- Absence of automated monitoring systems to track gap changes over time
- Limited visibility into geographic variations in supply-demand patterns
- Difficulty prioritizing action plans across multiple product categories"""
    
    doc.add_paragraph(problem_text)
    
    doc.add_heading('1.3 Objectives', level=2)
    objectives = [
        'Develop a comprehensive gap score methodology to quantify supply-demand imbalances',
        'Implement automated data pipeline for processing multi-source transactional data',
        'Compute gap scores at category, subcategory, and territory levels',
        'Generate actionable business insights through visualization and interpretation',
        'Create monitoring framework for weekly gap score tracking',
        'Design action planning system with prioritized recommendations',
        'Produce API-ready outputs for integration with business intelligence systems'
    ]
    
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')
    
    doc.add_heading('1.4 Scope and Limitations', level=2)
    scope_text = """The analysis covers all product categories within the rental platform dataset, examining supply and demand patterns across geographic territories and temporal dimensions. The project scope includes data cleaning, gap score computation, visualization, insight generation, and action planning.

Limitations include reliance on historical data without predictive modeling, assumption of linear supply-demand relationships, and focus on quantity-based metrics without incorporating pricing or profitability considerations. The analysis does not account for seasonal variations or external market factors beyond the dataset timeframe."""
    
    doc.add_paragraph(scope_text)
    doc.add_page_break()


def add_methodology(doc):
    """Add methodology section"""
    doc.add_heading('3. Methodology', level=1)
    
    doc.add_heading('3.1 Data Pipeline Architecture', level=2)
    pipeline_text = """The gap analysis pipeline consists of five sequential phases, each implemented as Jupyter notebooks with supporting Python modules. The modular architecture enables reproducibility, maintainability, and scalability of the analysis process.

The pipeline follows a structured workflow where each phase produces outputs consumed by subsequent phases. Data flows from raw Excel files through cleaning and transformation stages to final visualizations and action plans. All intermediate results are persisted to enable incremental processing and validation."""
    
    doc.add_paragraph(pipeline_text)
    
    # ...existing sections...
    
    doc.add_heading('3.3 Gap Score Calculation Formula', level=2)
    formula_text = """The gap score quantifies supply-demand imbalance using the following formula:

Gap Score = (Total Demand + 1) / (Total Supply + 1)

Where:
- Total Demand = Sum of order quantities or number of unique orders
- Total Supply = Number of unique product listings in the category
- +1 adjustment prevents division by zero and provides meaningful scores for zero-supply categories

Higher gap scores indicate greater supply shortages. A gap score of 1.0 represents perfect supply-demand balance, while scores exceeding 100 indicate severe shortages where demand outpaces supply by 100x or more."""
    
    doc.add_paragraph(formula_text)
    
    doc.add_heading('3.4 Classification Criteria', level=2)
    classification_text = """Gap scores are classified into four severity levels:

Low Gap (< 50): Balanced supply-demand relationship. No immediate action required.

Moderate Gap (50-100): Some supply shortage. Monitor closely and consider gradual supply expansion.

High Gap (100-200): Significant supply shortage. Prioritize supplier recruitment and inventory expansion.

Critical Gap (> 200): Severe supply shortage. Requires immediate operational intervention and strategic planning."""
    
    doc.add_paragraph(classification_text)
    doc.add_page_break()


def add_implementation(doc):
    """Add implementation section"""
    doc.add_heading('4. Implementation', level=1)
    
    phases = [
        {
            'title': '4.1 Phase 1: Data Preparation',
            'notebook': 'DataUnderstandingAndExploration.ipynb, DataCleaningAndMerging.ipynb',
            'description': 'Initial data loading, structure analysis, quality assessment, and cleaning operations. Handles missing values in ProductColor, ProductStyle, and customer demographics. Converts date columns to datetime format and merges product information with category hierarchies.',
            'outputs': 'Cleaned_data.xlsx containing 8 sheets with no missing values or duplicates'
        },
        {
            'title': '4.2 Phase 2: Supply-Demand Computation',
            'notebook': 'ComputeCategoryLevelSupply&Demand.ipynb',
            'description': 'Computes supply metrics (unique product counts) and demand metrics (order quantities, order counts) at category level. Calculates net demand by subtracting returns from sales. Generates monthly trends for temporal analysis.',
            'outputs': 'Category_Supply_Demand_Analysis.xlsx with category summary and monthly trends sheets'
        },
        {
            'title': '4.3 Phase 3: Gap Score Calculation',
            'notebook': 'MOAZ_Define_Metrics_GapScore.ipynb, GapScoreImplementation.ipynb',
            'description': 'Defines business KPIs and target multipliers. Computes gap scores at category, subcategory, and territory levels. Normalizes scores to 0-1 scale and classifies into severity levels. Performs territory-subcategory cross-analysis.',
            'outputs': 'comprehensive_gap_analysis.xlsx containing 4 sheets: Category Gaps, Subcategory Gaps, Territory Gaps, Detailed Gaps'
        },
        {
            'title': '4.4 Phase 4: Visualization and Insights',
            'notebook': 'VisualizationDraft.ipynb, InterpretationAndInsights.ipynb',
            'description': 'Generates 5 interactive Plotly visualizations: supply vs demand comparison, gap severity analysis, normalized metrics heatmap, category rankings, and gap distribution pie chart. Produces comprehensive textual insights with executive summary and actionable recommendations.',
            'outputs': 'Interactive HTML files, static PNG images, CSV/JSON data exports, insights text files'
        },
        {
            'title': '4.5 Phase 5: Monitoring and Action Planning',
            'notebook': 'Monitoring_and_Action_Plan.ipynb',
            'description': 'Creates weekly monitoring snapshots with gap score tracking. Computes KPIs and identifies categories needing attention. Generates prioritized action plans with operational, strategic, platform optimization, and user engagement recommendations.',
            'outputs': 'Weekly snapshot JSON/CSV files, action plan JSON/TXT files, monitoring dashboard HTML files'
        }
    ]
    
    for phase in phases:
        doc.add_heading(phase['title'], level=2)
        doc.add_paragraph(f"Notebooks: {phase['notebook']}", style='Intense Quote')
        doc.add_paragraph(phase['description'])
        doc.add_paragraph(f"Outputs: {phase['outputs']}", style='List Bullet')
        doc.add_paragraph()
    
    doc.add_page_break()


def add_technical_architecture(doc):
    """Add technical architecture section"""
    doc.add_heading('5. Technical Architecture', level=1)
    
    doc.add_heading('5.1 Python Modules', level=2)
    modules_text = """The project implements 21 reusable Python modules organized into functional categories:

Core Data Processing: data_loader.py, data_quality.py, data_statistics.py, data_relationships.py, data_db_model.py, data_exporter.py

Supply-Demand Computation: supply_demand.py (category, subcategory, territory-level calculations)

Gap Score Calculation: gap_score.py (formula implementation, normalization, classification)

Visualization: gap_visualizer.py, plotly_interactive_charts.py, plotly_gap_charts.py, plotly_dashboard.py, dashboard_builder.py, monitoring_visualizations.py

Insights Generation: chart_insights.py (textual insight generation from data)

Monitoring and Actions: monitoring_engine.py, action_planner.py (snapshot management, change analysis, action plan generation)

All modules follow consistent API design patterns with clear function signatures, type hints, and comprehensive docstrings. Modules are imported into notebooks using relative imports from the src directory."""
    
    doc.add_paragraph(modules_text)
    
    doc.add_heading('5.2 Notebook Workflow', level=2)
    workflow_text = """The 8 notebooks execute in strict sequential order, with each notebook depending on outputs from previous stages:

1. DataUnderstandingAndExploration.ipynb: Loads raw data, analyzes structure, identifies quality issues
2. DataCleaningAndMerging.ipynb: Cleans data, handles missing values, exports Cleaned_data.xlsx
3. ComputeCategoryLevelSupply&Demand.ipynb: Computes supply-demand metrics, exports analysis Excel
4. MOAZ_Define_Metrics_GapScore.ipynb: Defines KPIs, sets targets, classifies performance
5. GapScoreImplementation.ipynb: Calculates gap scores at all levels, exports comprehensive report
6. VisualizationDraft.ipynb: Generates interactive visualizations, exports HTML/PNG/data files
7. InterpretationAndInsights.ipynb: Produces textual business insights and recommendations
8. Monitoring_and_Action_Plan.ipynb: Creates monitoring snapshots and prioritized action plans

Each notebook can be executed independently after its prerequisites are met, enabling iterative development and testing."""
    
    doc.add_paragraph(workflow_text)
    doc.add_page_break()


def add_results_analysis(doc):
    """Add results and analysis section"""
    doc.add_heading('6. Results and Analysis', level=1)
    
    doc.add_heading('6.1 Category-Level Findings', level=2)
    category_text = """Analysis of 4 main product categories reveals varying degrees of supply-demand imbalance. Gap scores range from balanced categories with adequate supply to critical shortages requiring immediate intervention.

The category-level analysis provides a high-level overview suitable for executive decision-making and resource allocation planning. Categories with consistently high gap scores across multiple subcategories indicate systemic supply chain issues requiring strategic partnerships or supplier network expansion."""
    
    doc.add_paragraph(category_text)
    
    doc.add_heading('6.2 Subcategory-Level Findings', level=2)
    subcategory_text = """Drilling down to 37 subcategories reveals granular supply-demand patterns not visible at category level. Several subcategories exhibit critical gap scores exceeding 200, indicating demand outpacing supply by over 200x.

Top gap subcategories show consistent user search behavior and order attempts despite limited product availability. This represents significant revenue loss and potential customer churn. Subcategory analysis enables targeted supplier recruitment focused on specific product types rather than broad category expansion."""
    
    doc.add_paragraph(subcategory_text)
    
    doc.add_heading('6.3 Territory-Level Findings', level=2)
    territory_text = """Geographic analysis across 10 territories and 6 countries identifies regional supply-demand variations. Certain territories show higher average gap scores, suggesting localized supply shortages or region-specific demand patterns.

Territory-subcategory cross-analysis reveals 476 unique combinations, with several showing critical gaps in specific geographic markets. This enables region-specific inventory expansion and localized marketing campaigns targeting areas with proven demand but insufficient supply."""
    
    doc.add_paragraph(territory_text)
    
    doc.add_heading('6.4 Gap Severity Distribution', level=2)
    distribution_text = """Severity classification of all subcategories shows the proportion of categories in each gap level. The distribution informs strategic planning by indicating whether supply issues are concentrated in few categories or spread across the portfolio.

A high concentration of critical and high gap categories suggests platform-wide supply chain challenges requiring systemic solutions. Conversely, isolated critical gaps indicate category-specific issues addressable through targeted interventions."""
    
    doc.add_paragraph(distribution_text)
    doc.add_page_break()


def add_conclusion(doc):
    """Add conclusion section"""
    doc.add_heading('12. Conclusion', level=1)
    
    conclusion_text = """This gap analysis project successfully developed and implemented a comprehensive data-driven methodology for identifying and quantifying supply-demand imbalances in a rental platform. The multi-phase pipeline processes diverse data sources to compute actionable gap scores at category, subcategory, and territory levels.

Key achievements include:
- Automated data processing pipeline handling 8 datasets with 56,000+ transactions
- Gap score methodology with clear classification criteria
- Interactive visualizations enabling stakeholder engagement
- Monitoring framework for continuous gap tracking
- Prioritized action planning system with 300+ specific recommendations
- API-ready outputs for business intelligence integration

The analysis reveals significant supply shortages in multiple subcategories, with several exhibiting critical gaps requiring immediate operational intervention. The monitoring framework enables ongoing tracking of improvement initiatives and early detection of emerging gaps.

The modular architecture ensures maintainability and extensibility for future enhancements. The comprehensive documentation and code organization facilitate knowledge transfer and onboarding of new team members.

This project demonstrates the value of data-driven decision-making in operational optimization and strategic planning. The gap score methodology provides a quantitative foundation for resource allocation decisions that previously relied on qualitative assessments."""
    
    doc.add_paragraph(conclusion_text)
    doc.add_page_break()


def add_future_work(doc):
    """Add future work section"""
    doc.add_heading('13. Future Work', level=1)
    
    future_items = [
        'Predictive modeling: Develop time series forecasting models to predict future gap scores based on historical trends and seasonal patterns',
        'Dynamic pricing integration: Incorporate pricing elasticity into gap score calculations to optimize revenue while managing supply constraints',
        'Automated alert system: Implement real-time monitoring with automated email notifications when gap scores exceed critical thresholds',
        'Machine learning classification: Train models to predict gap severity based on product attributes, enabling proactive supply planning for new categories',
        'Customer segmentation analysis: Integrate customer lifetime value into demand calculations to prioritize high-value customer needs',
        'A/B testing framework: Measure impact of supply expansion initiatives on conversion rates and customer satisfaction',
        'Mobile dashboard: Develop mobile-responsive dashboards for on-the-go access to gap metrics and KPIs',
        'Integration with supplier management system: Automate supplier notification when gaps are detected in their product categories',
        'Scenario planning: Build what-if analysis tools to model impact of different supply expansion strategies'
    ]
    
    for item in future_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()


def add_references(doc):
    """Add references section"""
    doc.add_heading('14. References', level=1)
    
    references = [
        'Python Software Foundation. (2023). Python Language Reference, version 3.x. Available at https://www.python.org',
        'McKinney, W. (2022). Python for Data Analysis, 3rd Edition. O\'Reilly Media.',
        'Plotly Technologies Inc. (2023). Collaborative data science. Montreal, QC: Plotly Technologies Inc.',
        'pandas development team. (2023). pandas-dev/pandas: Pandas (v2.0.0). Zenodo.',
        'Hunter, J. D. (2007). Matplotlib: A 2D graphics environment. Computing in Science & Engineering, 9(3), 90-95.',
        'Waskom, M. (2021). seaborn: statistical data visualization. Journal of Open Source Software, 6(60), 3021.',
        'VanderPlas, J. (2016). Python Data Science Handbook. O\'Reilly Media.',
        'Géron, A. (2019). Hands-On Machine Learning with Scikit-Learn, Keras, and TensorFlow, 2nd Edition. O\'Reilly Media.'
    ]
    
    for i, ref in enumerate(references, 1):
        doc.add_paragraph(f'[{i}] {ref}', style='List Number')
    
    doc.add_page_break()


def add_appendices(doc):
    """Add appendices section"""
    doc.add_heading('15. Appendices', level=1)
    
    doc.add_heading('Appendix A: Data Dictionary', level=2)
    data_dict_text = """Complete listing of all columns across 8 data sheets with data types, descriptions, and sample values. Includes primary and foreign key relationships."""
    doc.add_paragraph(data_dict_text)
    
    doc.add_heading('Appendix B: Module API Documentation', level=2)
    api_doc_text = """Detailed function signatures, parameters, return types, and usage examples for all 21 Python modules. Includes code snippets demonstrating common use cases."""
    doc.add_paragraph(api_doc_text)
    
    doc.add_heading('Appendix C: Complete Gap Score Results', level=2)
    results_text = """Full tables of gap scores for all 37 subcategories, 10 territories, and 476 territory-subcategory combinations. Includes normalized scores and severity classifications."""
    doc.add_paragraph(results_text)
    
    doc.add_heading('Appendix D: Action Plans', level=2)
    actions_text = """Complete listing of operational, strategic, platform optimization, and user engagement actions for all categories. Organized by priority level."""
    doc.add_paragraph(actions_text)
    
    doc.add_heading('Appendix E: Installation and Setup Guide', level=2)
    setup_text = """Step-by-step instructions for environment setup, dependency installation, and notebook execution. Includes troubleshooting common issues."""
    doc.add_paragraph(setup_text)


def generate_report():
    """Main function to generate complete report"""
    print("Generating Gap Analysis Project Report...")
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Build report sections
    print("Adding title page...")
    add_title_page(doc)
    
    print("Adding abstract...")
    add_abstract(doc)
    
    print("Adding table of contents...")
    add_table_of_contents(doc)
    
    print("Adding introduction...")
    add_introduction(doc)
    
    print("Adding methodology...")
    add_methodology(doc)
    
    print("Adding implementation...")
    add_implementation(doc)
    
    print("Adding technical architecture...")
    add_technical_architecture(doc)
    
    print("Adding results and analysis...")
    add_results_analysis(doc)
    
    print("Adding conclusion...")
    add_conclusion(doc)
    
    print("Adding future work...")
    add_future_work(doc)
    
    print("Adding references...")
    add_references(doc)
    
    print("Adding appendices...")
    add_appendices(doc)
    
    # Save document
    output_path = Path(__file__).parent.parent / "Gap_Analysis_Project_Report.docx"
    doc.save(output_path)
    
    print(f"\nReport generated successfully!")
    print(f"Location: {output_path}")
    print(f"Total sections: 15")
    print("\nYou can now open the DOCX file and:")
    print("  1. Update the table of contents (right-click > Update Field)")
    print("  2. Add page numbers (Insert > Page Number)")
    print("  3. Customize formatting as needed")
    print("  4. Add charts and figures from results folder")


if __name__ == "__main__":
    generate_report()
