"""
Gap Analysis Project - Comprehensive Report Generator (v2)
Generates a detailed technical report explaining the project architecture, workflow, and results.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
from pathlib import Path

def setup_document_styles(doc):
    """Configure document styles for a professional look."""
    # Normal text
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Headings
    for i in range(1, 4):
        style = doc.styles[f'Heading {i}']
        font = style.font
        font.name = 'Calibri Light'
        font.color.rgb = RGBColor(44, 62, 80)  # Dark Blue-Grey
        if i == 1:
            font.size = Pt(16)
            font.bold = True
        elif i == 2:
            font.size = Pt(14)
            font.bold = True
        elif i == 3:
            font.size = Pt(12)
            font.bold = True

def add_title_page(doc):
    """Add a clean, professional title page."""
    # Spacer
    for _ in range(5):
        doc.add_paragraph()
        
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Supply-Demand Gap Analysis\nTechnical Report')
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)  # Navy Blue
    
    doc.add_paragraph()
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Rently 2026 Platform Optimization')
    run.font.size = Pt(16)
    run.font.italic = True
    
    for _ in range(4):
        doc.add_paragraph()
        
    # Info Block
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('Prepared for:\nRently 2026 Stakeholders\n\nPrepared by:\nGap Analysis Data Team\n\nDate: ' + datetime.now().strftime("%B %d, %Y"))
    run.font.size = Pt(12)
    
    doc.add_page_break()

def add_introduction(doc):
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        "This report details the technical implementation and results of the Supply-Demand Gap Analysis project for the Rently 2026 platform. "
        "The primary objective is to identify product categories experiencing significant supply shortages relative to customer demand, enabling data-driven inventory optimization."
    )
    doc.add_paragraph(
        "The system utilizes a modular Python architecture to process transactional data, compute gap scores, and generate actionable insights through interactive visualizations and automated monitoring reports."
    )

def add_system_architecture(doc):
    doc.add_heading('2. System Architecture & Integration', level=1)
    
    doc.add_heading('2.1 The "One Unit" Concept', level=2)
    doc.add_paragraph(
        "The project is designed as a cohesive unit where Jupyter Notebooks serve as the orchestration layer, and the `src` directory contains the core business logic. "
        "This separation of concerns ensures that the analysis is reproducible, modular, and easily maintainable."
    )
    
    doc.add_paragraph(
        "• **Orchestration (Notebooks)**: Define the step-by-step workflow, from data loading to insight generation. They call functions from the `src` modules to perform heavy lifting.\n"
        "• **Core Logic (src/)**: Contains reusable Python modules for data processing, statistical analysis, and visualization. This allows the same logic to be used across multiple notebooks and potentially in a production API."
    )
    
    doc.add_heading('2.2 Integration Flow', level=2)
    doc.add_paragraph(
        "Data flows sequentially through the pipeline:\n"
        "1. **Raw Data** is ingested and cleaned using `data_loader` and `data_quality` modules.\n"
        "2. **Supply & Demand Metrics** are computed using `supply_demand.py`.\n"
        "3. **Gap Scores** are calculated via `gap_score.py`.\n"
        "4. **Visualizations** are generated using `plotly_interactive_charts.py`.\n"
        "5. **Monitoring** is handled by `monitoring_engine.py`."
    )

def add_detailed_workflow(doc):
    doc.add_heading('3. Detailed Workflow Analysis', level=1)
    doc.add_paragraph("The analysis is executed through a sequence of 8 integrated notebooks:")
    
    notebooks = [
        ("1. DataUnderstandingAndExploration.ipynb", 
         "Exploratory Data Analysis (EDA)", 
         "Loads raw Excel sheets, analyzes schema relationships, and identifies data quality issues (missing values in `ProductColor`, `ProductStyle`). Uses `data_relationships.py` to infer foreign keys."),
        
        ("2. DataCleaningAndMerging.ipynb", 
         "Data Preprocessing", 
         "Implements cleaning strategies: fills missing values (e.g., 'Mixed' for color), standardizes formats, and merges disparate datasets into a unified `Cleaned_data.xlsx`. This file serves as the single source of truth for downstream analysis."),
        
        ("3. ComputeCategoryLevelSupply&Demand.ipynb", 
         "Metric Computation", 
         "Aggregates data to calculate fundamental metrics: Total Sales (Demand) and Unique Product Listings (Supply) at the category level. Outputs `Category_Supply_Demand_Analysis.xlsx`."),
        
        ("4. MOAZ_Define_Metrics_GapScore.ipynb", 
         "KPI Definition", 
         "Establishes the business logic for success. Defines target multipliers (e.g., 1.1x revenue target) and configures the scoring parameters in `config.py`."),
        
        ("5. GapScoreImplementation.ipynb", 
         "Core Analysis", 
         "The heart of the system. Implements the Gap Score formula: `(Demand + 1) / (Supply + 1)`. Computes scores for Categories, Subcategories, and Territories. Classifies gaps as Low, Moderate, High, or Critical."),
        
        ("6. VisualizationDraft.ipynb", 
         "Visual Reporting", 
         "Generates the suite of interactive Plotly charts found in `results/visualizations`. Creates the HTML dashboards that allow stakeholders to explore the data dynamically."),
        
        ("7. InterpretationAndInsights.ipynb", 
         "Automated Insights", 
         "Uses `chart_insights.py` to programmatically analyze the data and generate textual explanations, identifying top performing and underperforming areas."),
        
        ("8. Monitoring_and_Action_Plan.ipynb", 
         "Continuous Improvement", 
         "Simulates a weekly monitoring cycle. Compares current data against historical snapshots to detect trends. Generates `action_plan_week_X.json` with specific recommendations.")
    ]
    
    for title, purpose, desc in notebooks:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        doc.add_paragraph(f"Purpose: {purpose}", style='Intense Quote')
        doc.add_paragraph(desc)
        doc.add_paragraph()

def add_outputs_breakdown(doc):
    doc.add_heading('4. Outputs & Results Explained', level=1)
    doc.add_paragraph("The `results` directory contains structured outputs designed for different stakeholders:")
    
    doc.add_heading('4.1 Data Exports (results/data)', level=2)
    doc.add_paragraph(
        "• **gap_summary.json/csv**: Lightweight summaries of the gap analysis, ideal for frontend applications or quick lookups.\n"
        "• **Purpose**: Provides immediate access to key metrics without parsing large Excel files."
    )
    
    doc.add_heading('4.2 Comprehensive Reports (results/reports)', level=2)
    doc.add_paragraph(
        "• **comprehensive_gap_analysis.xlsx**: The master report containing detailed sheets for Categories, Subcategories, and Territories.\n"
        "• **Purpose**: Deep-dive analysis for data analysts and business strategists."
    )
    
    doc.add_heading('4.3 Interactive Visualizations (results/visualizations)', level=2)
    doc.add_paragraph(
        "• **HTML Dashboards**: Interactive files (e.g., `01_supply_vs_demand_subcategory.html`) that allow users to hover, zoom, and filter data.\n"
        "• **Purpose**: Presentation tools for stakeholder meetings."
    )
    
    doc.add_heading('4.4 Monitoring Artifacts (results/monitoring_reports)', level=2)
    doc.add_paragraph(
        "• **Snapshots (JSON/CSV)**: Point-in-time records of gap scores.\n"
        "• **Action Plans (JSON/TXT)**: Prioritized lists of recommended actions (e.g., 'Increase supply for Projectors').\n"
        "• **Purpose**: Operational tracking to ensure gaps are being closed over time."
    )

def add_technical_details(doc):
    doc.add_heading('5. Technical Methodology', level=1)
    
    doc.add_heading('5.1 Gap Score Formula', level=2)
    doc.add_paragraph("The Gap Score is calculated as:")
    doc.add_paragraph("Gap Score = (Total Quantity Sold + 1) / (Unique Products + 1)", style='Quote')
    doc.add_paragraph(
        "The `+1` smoothing term prevents division by zero and stabilizes scores for categories with very low volume. "
        "A higher score indicates a larger disparity between demand and supply."
    )
    
    doc.add_heading('5.2 Severity Classification', level=2)
    doc.add_paragraph(
        "Scores are classified into four actionable levels:\n"
        "• **Low (< 50)**: Healthy balance.\n"
        "• **Moderate (50-100)**: Monitor closely.\n"
        "• **High (100-200)**: Needs attention.\n"
        "• **Critical (> 200)**: Immediate action required."
    )

def add_conclusion(doc):
    doc.add_heading('6. Conclusion', level=1)
    doc.add_paragraph(
        "The Rently 2026 Gap Analysis project provides a robust, automated framework for supply chain optimization. "
        "By integrating data processing, advanced analytics, and automated reporting, the system transforms raw transaction data into strategic assets. "
        "The modular architecture ensures that the system can scale with the business, accommodating new data sources and evolving metric definitions."
    )

def generate_report():
    doc = Document()
    setup_document_styles(doc)
    
    add_title_page(doc)
    add_introduction(doc)
    add_system_architecture(doc)
    add_detailed_workflow(doc)
    add_outputs_breakdown(doc)
    add_technical_details(doc)
    add_conclusion(doc)
    
    output_path = Path(__file__).parent.parent / "Rently_Gap_Analysis_Comprehensive_Report.docx"
    doc.save(output_path)
    print(f"Report generated at: {output_path}")

if __name__ == "__main__":
    generate_report()
